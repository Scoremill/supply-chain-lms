#!/usr/bin/env python3
"""
Generate the Strategem "Build A Home" LMS User Manual as a Word document.
Usage: python3 scripts/generate-user-manual.py
Output: Strategem_Build_A_Home_User_Manual.docx
"""

import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Branding constants ──────────────────────────────────────────────────────
ORANGE = RGBColor(0xF9, 0x73, 0x16)       # #f97316
DARK_ORANGE = RGBColor(0xEA, 0x58, 0x0C)  # #ea580c
DARK_BLUE = RGBColor(0x1E, 0x29, 0x3B)    # Dark heading color
GRAY_600 = RGBColor(0x4B, 0x55, 0x63)
GRAY_500 = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
LIGHT_GRAY_BG = RGBColor(0xF3, 0xF4, 0xF6)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
LOGO_PATH = os.path.join(PROJECT_ROOT, "public", "stratagem-logo-full.png")
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "Strategem_Build_A_Home_User_Manual.docx")


# ── Helper functions ─────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_screenshot_placeholder(doc, caption, width_inches=5.5):
    """Add a bordered gray placeholder box for a screenshot."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")

    # Set cell borders
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="8" w:space="0" w:color="D1D5DB"/>'
        '  <w:left w:val="single" w:sz="8" w:space="0" w:color="D1D5DB"/>'
        '  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="D1D5DB"/>'
        '  <w:right w:val="single" w:sz="8" w:space="0" w:color="D1D5DB"/>'
        '</w:tcBorders>'
    )
    tc_pr.append(borders)

    # Set cell width
    tc_w = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width_inches * 1440)}" w:type="dxa"/>')
    tc_pr.append(tc_w)

    # Add placeholder text
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(30)
    p.space_after = Pt(30)

    run = p.add_run(f"\n[Screenshot: {caption}]\n")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = GRAY_500

    # Add caption below
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.space_after = Pt(12)
    run = caption_p.add_run(f"Figure: {caption}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GRAY_500


def add_callout_box(doc, label, text, box_type="tip"):
    """Add a formatted callout box (Tip, Note, Important)."""
    color_map = {
        "tip":       ("10B981", "ECFDF5", "065F46"),
        "note":      ("3B82F6", "EFF6FF", "1E3A5F"),
        "important": ("F97316", "FFF7ED", "9A3412"),
        "warning":   ("EF4444", "FEF2F2", "991B1B"),
    }
    border_hex, bg_hex, text_hex = color_map.get(box_type, color_map["tip"])

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg_hex)

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>'
        f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="{border_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>'
        '</w:tcBorders>'
    )
    tc_pr.append(borders)

    p = cell.paragraphs[0]
    p.space_before = Pt(6)
    p.space_after = Pt(6)

    label_run = p.add_run(f"{label.upper()}: ")
    label_run.font.bold = True
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = RGBColor(
        int(text_hex[0:2], 16), int(text_hex[2:4], 16), int(text_hex[4:6], 16)
    )

    text_run = p.add_run(text)
    text_run.font.size = Pt(10)
    text_run.font.color.rgb = RGBColor(
        int(text_hex[0:2], 16), int(text_hex[2:4], 16), int(text_hex[4:6], 16)
    )

    doc.add_paragraph()  # spacer


def add_body(doc, text):
    """Add a body paragraph with consistent formatting."""
    p = doc.add_paragraph()
    p.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    return p


def add_body_bold_lead(doc, bold_part, rest):
    """Add a paragraph with a bold lead-in followed by normal text."""
    p = doc.add_paragraph()
    p.space_after = Pt(8)
    b = p.add_run(bold_part)
    b.font.bold = True
    b.font.size = Pt(11)
    b.font.color.rgb = DARK_BLUE
    r = p.add_run(rest)
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_BLUE
    return p


def add_numbered_step(doc, number, text):
    """Add a numbered step paragraph."""
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    num = p.add_run(f"Step {number}: ")
    num.font.bold = True
    num.font.size = Pt(11)
    num.font.color.rgb = ORANGE
    body = p.add_run(text)
    body.font.size = Pt(11)
    body.font.color.rgb = DARK_BLUE
    return p


def add_bullet(doc, text):
    """Add a bullet-point list item."""
    p = doc.add_paragraph(style="List Bullet")
    p.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    return p


def add_section_break(doc):
    """Add a page break."""
    doc.add_page_break()


def style_heading(heading, level=1):
    """Apply Strategem orange styling to headings."""
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = ORANGE
            run.font.size = Pt(24)
        elif level == 2:
            run.font.color.rgb = DARK_ORANGE
            run.font.size = Pt(18)
        elif level == 3:
            run.font.color.rgb = DARK_BLUE
            run.font.size = Pt(14)


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    style_heading(h, 1)
    return h


def h2(doc, text):
    h = doc.add_heading(text, level=2)
    style_heading(h, 2)
    return h


def h3(doc, text):
    h = doc.add_heading(text, level=3)
    style_heading(h, 3)
    return h


def add_role_table(doc):
    """Add the user roles summary table."""
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Role", "Description", "Key Capabilities"]
    rows_data = [
        ["Student", "End-user learner enrolled in the course",
         "View modules, watch videos, take quizzes, earn certificate"],
        ["Company Administrator", "Manages students within their organization",
         "View company stats, monitor student progress, export data, view takeaways"],
        ["Super Administrator", "Full platform access and control",
         "All Company Admin capabilities plus: manage users, companies, modules, codes, analytics, emulation"],
    ]

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, "F97316")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)

    for r, row_data in enumerate(rows_data):
        for c, value in enumerate(row_data):
            cell = table.cell(r + 1, c)
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = DARK_BLUE
            if r % 2 == 1:
                set_cell_shading(cell, "FFF7ED")

    doc.add_paragraph()  # spacer


# ── Main document builder ────────────────────────────────────────────────────

def build_manual():
    doc = Document()

    # ── Default font styling ─────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_BLUE

    # Set default paragraph spacing
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(8)
    paragraph_format.line_spacing = 1.15

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ═══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════

    # Add some vertical space
    for _ in range(4):
        doc.add_paragraph()

    # Logo
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(3.5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Strategem Logo]")
        run.font.size = Pt(18)
        run.font.color.rgb = ORANGE

    doc.add_paragraph()  # spacer

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Build A Home")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = ORANGE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Learning Management System")
    run.font.size = Pt(24)
    run.font.color.rgb = DARK_BLUE

    doc.add_paragraph()

    manual_title = doc.add_paragraph()
    manual_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = manual_title.add_run("User Manual")
    run.font.size = Pt(20)
    run.font.color.rgb = GRAY_600

    doc.add_paragraph()
    doc.add_paragraph()

    # Version and date
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Version 1.0  |  {date.today().strftime('%B %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY_500

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run("Strategem LLC  |  Confidential")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY_500

    add_section_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════════════

    h1(doc, "Table of Contents")
    doc.add_paragraph()

    toc_items = [
        ("Part 1:", "General Information"),
        ("  1.1", "About This Manual"),
        ("  1.2", "Platform Overview"),
        ("  1.3", "System Requirements"),
        ("", ""),
        ("Part 2:", "Student User Guide"),
        ("  2.1", "Getting Started"),
        ("  2.2", "Your Dashboard"),
        ("  2.3", "Learning Modules"),
        ("  2.4", "Quizzes & Assessments"),
        ("  2.5", "Key Takeaways"),
        ("  2.6", "Chat To A Builder \u2014 AI Assistant"),
        ("  2.7", "Course Completion & Certification"),
        ("  2.8", "End-of-Course Survey"),
        ("  2.9", "Getting Help"),
        ("", ""),
        ("Part 3:", "Company Administrator Guide"),
        ("  3.1", "Getting Started"),
        ("  3.2", "Company Admin Dashboard"),
        ("  3.3", "Managing Your Students"),
        ("  3.4", "Viewing Student Takeaways"),
        ("  3.5", "Company Codes"),
        ("  3.6", "Exporting Data"),
        ("  3.7", "Getting Help"),
        ("", ""),
        ("Part 4:", "Super Administrator Guide"),
        ("  4.1", "Getting Started"),
        ("  4.2", "Platform Analytics"),
        ("  4.3", "User Management"),
        ("  4.4", "Company Management"),
        ("  4.5", "Company Code Management"),
        ("  4.6", "Viewing User Takeaways"),
        ("  4.7", "Module Management"),
        ("  4.8", "Survey Results"),
        ("  4.9", "Data Exports"),
        ("  4.10", "User Emulation"),
        ("  4.11", "Getting Help"),
        ("", ""),
        ("Part 5:", "Troubleshooting & FAQ"),
        ("", ""),
        ("Appendix:", "Glossary of Terms"),
    ]

    for num, title_text in toc_items:
        if not num and not title_text:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        if num.startswith("Part") or num.startswith("Appendix"):
            r1 = p.add_run(f"{num} ")
            r1.font.bold = True
            r1.font.size = Pt(12)
            r1.font.color.rgb = ORANGE
            r2 = p.add_run(title_text)
            r2.font.bold = True
            r2.font.size = Pt(12)
            r2.font.color.rgb = DARK_BLUE
        else:
            r1 = p.add_run(f"{num}  ")
            r1.font.size = Pt(11)
            r1.font.color.rgb = GRAY_500
            r2 = p.add_run(title_text)
            r2.font.size = Pt(11)
            r2.font.color.rgb = DARK_BLUE

    add_section_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PART 1: GENERAL INFORMATION
    # ═══════════════════════════════════════════════════════════════════════

    h1(doc, "Part 1: General Information")
    doc.add_paragraph()

    # ── 1.1 About This Manual ──
    h2(doc, "1.1  About This Manual")

    add_body(doc,
        "This document serves as the official user manual for the Strategem "
        "\"Build A Home\" Learning Management System (LMS). It provides comprehensive "
        "guidance for all platform users, from students progressing through the course "
        "curriculum to administrators managing the platform and its users."
    )
    add_body(doc,
        "The manual is organized into role-based sections so that each user type can "
        "quickly locate the information most relevant to their responsibilities. Each "
        "section is designed to function as a standalone reference, meaning it can be "
        "distributed independently to users of that role without requiring the full document."
    )

    h3(doc, "How to Use This Manual")
    add_body(doc,
        "Depending on your role within the platform, you may find it most efficient to "
        "focus on the section that applies to you:"
    )
    add_bullet(doc, "Students should refer to Part 2: Student User Guide.")
    add_bullet(doc, "Company Administrators should refer to Part 3: Company Administrator Guide.")
    add_bullet(doc, "Super Administrators should refer to Part 4: Super Administrator Guide.")
    add_body(doc,
        "All users are encouraged to review Part 1 (General Information) for an overview "
        "of the platform and Part 5 (Troubleshooting & FAQ) for answers to common questions."
    )

    h3(doc, "Conventions Used in This Manual")
    add_body(doc, "Throughout this document, you will encounter the following conventions:")
    add_callout_box(doc, "Tip", "Provides helpful suggestions to enhance your experience on the platform.", "tip")
    add_callout_box(doc, "Note", "Highlights important information that you should be aware of.", "note")
    add_callout_box(doc, "Important", "Indicates critical information that may affect your ability to use the platform or complete the course.", "important")
    add_body(doc,
        "Gray placeholder boxes marked with \"[Screenshot: ...]\" indicate where "
        "screenshots will be inserted in future versions of this manual. These placeholders "
        "describe the content that should appear in the image."
    )

    # ── 1.2 Platform Overview ──
    h2(doc, "1.2  Platform Overview")

    add_body(doc,
        "The Strategem \"Build A Home\" LMS is a web-based training platform designed for "
        "the construction industry. It delivers a structured, 10-module curriculum that covers "
        "the essential knowledge required to successfully build a home. The course combines "
        "video instruction, rich written content, interactive knowledge checks, and formal "
        "quiz assessments to ensure a thorough understanding of each topic."
    )

    h3(doc, "Key Features")
    add_bullet(doc, "10 comprehensive learning modules covering all major aspects of home construction.")
    add_bullet(doc, "Video content integrated into each module for visual and auditory learning.")
    add_bullet(doc, "Interactive knowledge checks embedded within module content for self-assessment.")
    add_bullet(doc, "Formal quiz assessments at the end of each module with an 85% passing threshold.")
    add_bullet(doc, "Automatic certificate generation upon successful completion of all modules.")
    add_bullet(doc, "AI-powered \"Chat To A Builder\" assistant for real-time questions about course content.")
    add_bullet(doc, "Progress tracking and analytics for students and administrators.")
    add_bullet(doc, "Company-based enrollment and student management.")
    add_bullet(doc, "Multi-role access control (Student, Company Admin, Super Admin).")

    h3(doc, "User Roles")
    add_body(doc,
        "The platform supports three distinct user roles, each with different levels of "
        "access and functionality. The table below summarizes these roles:"
    )
    add_role_table(doc)

    # ── 1.3 System Requirements ──
    h2(doc, "1.3  System Requirements")

    add_body(doc,
        "The Build A Home LMS is a web-based application that runs entirely in your "
        "browser. No software installation is required. To ensure the best experience, "
        "please verify that your system meets the following requirements."
    )

    h3(doc, "Supported Browsers")
    add_bullet(doc, "Google Chrome (version 90 or later) \u2014 Recommended")
    add_bullet(doc, "Mozilla Firefox (version 90 or later)")
    add_bullet(doc, "Microsoft Edge (version 90 or later)")
    add_bullet(doc, "Apple Safari (version 14 or later)")

    h3(doc, "Internet Connection")
    add_body(doc,
        "A stable broadband internet connection is required. Video content is streamed "
        "directly within the platform, so a minimum download speed of 5 Mbps is "
        "recommended for smooth playback without buffering."
    )

    h3(doc, "Device Compatibility")
    add_body(doc,
        "The platform is designed to be responsive and can be accessed from desktop "
        "computers, laptops, and tablets. While the platform is accessible on mobile "
        "devices, a screen width of at least 768 pixels (tablet size or larger) is "
        "recommended for the best learning experience, particularly when viewing module "
        "content and taking quizzes."
    )

    add_callout_box(doc, "Tip",
        "For the best experience, use a desktop or laptop computer with Google Chrome "
        "and ensure your browser is up to date.",
        "tip"
    )

    add_section_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PART 2: STUDENT USER GUIDE
    # ═══════════════════════════════════════════════════════════════════════

    h1(doc, "Part 2: Student User Guide")
    add_body(doc,
        "This section of the manual is intended for students enrolled in the Build A Home "
        "course. It covers everything you need to know to access the platform, navigate "
        "the learning modules, complete assessments, and earn your certificate of completion."
    )

    # ── 2.1 Getting Started ──
    h2(doc, "2.1  Getting Started")

    h3(doc, "Accessing the Platform")
    add_body(doc,
        "To access the Build A Home LMS, open your web browser and navigate to the "
        "platform URL provided by your organization or administrator. If you have not "
        "received a URL, contact your company administrator or Strategem support for "
        "assistance."
    )
    add_screenshot_placeholder(doc, "Sign-In / Registration Page \u2014 Full View")

    h3(doc, "Watch the Introduction Video")
    add_callout_box(doc, "Important",
        "Before registering or beginning the course, it is strongly recommended that "
        "you watch the platform introduction video. This 2-minute video provides a "
        "concise overview of the course structure, explains how to navigate the platform, "
        "and sets expectations for your learning journey.",
        "important"
    )
    add_body(doc,
        "On the sign-in page, you will see a prominent orange button labeled "
        "\"Click for an Introduction to the course.\" This button features a pulsating "
        "animation to draw your attention. Click this button to open a modal window "
        "containing the introduction video. Use the standard video controls (play, pause, "
        "volume, full-screen) to watch the video at your own pace."
    )
    add_body(doc,
        "The video covers the following topics:"
    )
    add_bullet(doc, "An overview of the Build A Home course and its 10-module structure.")
    add_bullet(doc, "How to navigate through modules, videos, and written content.")
    add_bullet(doc, "What to expect from quizzes and how scoring works.")
    add_bullet(doc, "How to earn your certificate upon completion.")

    add_screenshot_placeholder(doc, "Introduction Video Button (Orange Pulsating Button on Sign-In Page)")
    add_screenshot_placeholder(doc, "Introduction Video Modal \u2014 Playing the Video")

    h3(doc, "Creating Your Account")
    add_body(doc,
        "If you are a new user, you will need to create an account before accessing the "
        "course content. Follow these steps to register:"
    )
    add_numbered_step(doc, 1,
        "On the sign-in page, click the \"Register\" tab at the top of the authentication "
        "card on the right side of the screen."
    )
    add_numbered_step(doc, 2,
        "Enter your full name in the \"Full Name\" field. Use the name you would like to "
        "appear on your certificate of completion."
    )
    add_numbered_step(doc, 3,
        "Enter your email address in the \"Email\" field. This will be your login credential "
        "and the address used for any account-related communications."
    )
    add_numbered_step(doc, 4,
        "If your employer or organization has provided you with a Company Code, enter it in "
        "the \"Company Code\" field. This code associates your account with your organization "
        "and allows your company administrator to track your progress. If you do not have a "
        "company code, you may leave this field blank for public access."
    )
    add_numbered_step(doc, 5,
        "Create a password of at least 6 characters and enter it in the \"Password\" field."
    )
    add_numbered_step(doc, 6,
        "Re-enter your password in the \"Confirm Password\" field to verify it."
    )
    add_numbered_step(doc, 7,
        "Click the \"Create Account\" button. If all information is valid, your account will "
        "be created and you will be automatically signed in and redirected to your dashboard."
    )

    add_callout_box(doc, "Note",
        "If you received a Company Code from your employer, it is important to enter it "
        "during registration. You cannot retroactively add a company code to your account. "
        "If you forget to enter it, contact your administrator for assistance.",
        "note"
    )

    add_screenshot_placeholder(doc, "Registration Form \u2014 Filled Out with Company Code")

    h3(doc, "Signing In")
    add_body(doc,
        "If you already have an account, sign in using the following steps:"
    )
    add_numbered_step(doc, 1,
        "On the sign-in page, ensure the \"Login\" tab is selected at the top of the "
        "authentication card."
    )
    add_numbered_step(doc, 2, "Enter the email address you used during registration.")
    add_numbered_step(doc, 3, "Enter your password.")
    add_numbered_step(doc, 4,
        "Click the \"Sign In\" button. Upon successful authentication, you will be "
        "redirected to your student dashboard."
    )

    add_screenshot_placeholder(doc, "Login Form")

    h3(doc, "Forgot Password / Reset Password")
    add_body(doc,
        "If you have forgotten your password, the platform provides a self-service "
        "password reset feature:"
    )
    add_numbered_step(doc, 1,
        "On the sign-in page, click the \"Forgot Password?\" link located beneath "
        "the password field."
    )
    add_numbered_step(doc, 2,
        "Enter the email address associated with your account and submit the form."
    )
    add_numbered_step(doc, 3,
        "Check your email inbox (and spam/junk folder) for a password reset link."
    )
    add_numbered_step(doc, 4,
        "Click the link in the email to open the password reset page. Enter and confirm "
        "your new password, then submit."
    )
    add_numbered_step(doc, 5,
        "After resetting your password, return to the sign-in page and log in with "
        "your new credentials."
    )

    add_callout_box(doc, "Tip",
        "Password reset links expire after a limited time. If your link has expired, "
        "simply request a new one by repeating the process above.",
        "tip"
    )

    # ── 2.2 Your Dashboard ──
    h2(doc, "2.2  Your Dashboard")

    add_body(doc,
        "After signing in, you will be taken to your student dashboard. The dashboard "
        "is your central hub for accessing course modules, tracking your progress, and "
        "viewing your overall completion status."
    )

    h3(doc, "Dashboard Layout")
    add_body(doc,
        "The student dashboard is organized into several key areas:"
    )
    add_body_bold_lead(doc, "Welcome Banner: ",
        "Displayed at the top of the dashboard, the welcome banner greets you by name "
        "and provides a summary of your current progress, including the number of modules "
        "completed and your overall course completion percentage."
    )
    add_body_bold_lead(doc, "Progress Statistics: ",
        "A set of statistic cards showing key metrics at a glance, including modules "
        "completed, quizzes passed, and your overall progress percentage."
    )
    add_body_bold_lead(doc, "Module Grid: ",
        "Below the statistics, you will see a grid of cards representing each of the "
        "10 course modules. Each card displays the module title, a brief description, "
        "and your current status for that module (Not Started, In Progress, or Completed). "
        "Modules that are locked (because the preceding module has not been completed) "
        "will be visually distinguished."
    )

    add_screenshot_placeholder(doc, "Student Dashboard \u2014 Overview with Progress Stats and Module Grid")

    h3(doc, "Understanding Your Progress Indicators")
    add_body(doc,
        "Each module card on the dashboard includes visual indicators to help you "
        "understand your progress at a glance:"
    )
    add_bullet(doc, "A progress bar showing the percentage of content viewed within the module.")
    add_bullet(doc, "A status badge indicating whether the module is Not Started, In Progress, or Completed.")
    add_bullet(doc, "A quiz score display (if a quiz has been attempted) showing your highest score.")
    add_bullet(doc, "A lock icon on modules that are not yet available (sequential progression is required).")

    add_callout_box(doc, "Note",
        "Modules must be completed in sequential order. You cannot skip ahead to a later "
        "module without first passing the quiz for the preceding module with a score of "
        "85% or higher.",
        "note"
    )

    # ── 2.3 Learning Modules ──
    h2(doc, "2.3  Learning Modules")

    add_body(doc,
        "The Build A Home course consists of 10 comprehensive learning modules. Each "
        "module follows a consistent structure designed to maximize your understanding "
        "and retention of the material."
    )

    h3(doc, "Module Structure")
    add_body(doc, "Every module includes the following components:")
    add_body_bold_lead(doc, "Video Content: ",
        "Each module begins with an instructional video that introduces the key concepts "
        "and topics covered in the module. Videos are embedded directly within the module "
        "page and include standard playback controls (play, pause, volume, seek, and "
        "full-screen)."
    )
    add_body_bold_lead(doc, "Written Content: ",
        "Following the video, each module contains detailed written material that expands "
        "on the topics introduced in the video. The written content includes formatted "
        "text, images, tables, and diagrams where applicable."
    )
    add_body_bold_lead(doc, "Knowledge Checks: ",
        "Embedded throughout the written content, you will encounter interactive "
        "\"Knowledge Check\" questions. These are self-assessment questions designed to "
        "reinforce your learning as you read. They do not count toward your quiz score."
    )
    add_body_bold_lead(doc, "Module Quiz: ",
        "At the end of each module, there is a formal assessment quiz. You must score "
        "85% or higher to pass the quiz and unlock the next module."
    )

    h3(doc, "Starting a Module")
    add_body(doc,
        "To begin a module, click on its card from your student dashboard. You will be "
        "taken to the module page, which displays the video and content in a structured "
        "layout."
    )
    add_screenshot_placeholder(doc, "Module Page \u2014 Video Player and Content Area")

    h3(doc, "Watching Video Content")
    add_body(doc,
        "The video player is located at the top of the module page. Click the play button "
        "to begin watching. You may pause, rewind, fast-forward, adjust volume, or "
        "switch to full-screen mode using the controls at the bottom of the player. "
        "Watching the video is an important part of the learning experience and covers "
        "material that may appear on the quiz."
    )

    add_callout_box(doc, "Tip",
        "Take notes while watching the video. The concepts presented will be reinforced "
        "in the written content below, and you may be tested on them in the module quiz.",
        "tip"
    )

    h3(doc, "Reading Module Content")
    add_body(doc,
        "Below the video, the module's written content is displayed in a scrollable area. "
        "The content is organized with headings, subheadings, and a Module Map (table of "
        "contents) that allows you to quickly navigate to specific sections within the module."
    )
    add_body_bold_lead(doc, "Module Map: ",
        "Located at the top of the content area, the Module Map provides a clickable "
        "list of all sections within the module. Click any section title to jump "
        "directly to that section."
    )
    add_body_bold_lead(doc, "Expandable Sections: ",
        "Some content is organized within collapsible sections indicated by an orange "
        "header bar. Click the bar (labeled \"CLICK TO EXPAND\") to reveal the content "
        "within. Click again to collapse it."
    )

    add_screenshot_placeholder(doc, "Module Content \u2014 Module Map and Scrollable Content Area")
    add_screenshot_placeholder(doc, "Expandable Content Section \u2014 Collapsed and Expanded States")

    h3(doc, "Knowledge Checks")
    add_body(doc,
        "As you read through the module content, you will encounter Knowledge Check "
        "cards. These are interactive, self-paced questions embedded directly in the "
        "content to help you test your understanding as you learn."
    )
    add_body(doc, "To use a Knowledge Check:")
    add_numbered_step(doc, 1, "Read the question presented in the Knowledge Check card.")
    add_numbered_step(doc, 2, "Select the answer you believe is correct from the available options.")
    add_numbered_step(doc, 3,
        "Click the \"Reveal Answer\" button to see whether your selection was correct."
    )
    add_numbered_step(doc, 4,
        "Review the correct answer and explanation. If you answered incorrectly, "
        "consider re-reading the relevant section before moving on."
    )

    add_callout_box(doc, "Note",
        "Knowledge Checks are for self-assessment only. They do not count toward your "
        "module quiz score. However, they cover the same material that appears on the "
        "quiz, so use them as a study tool.",
        "note"
    )

    add_screenshot_placeholder(doc, "Knowledge Check Card \u2014 Question with Answer Options")

    # ── 2.4 Quizzes & Assessments ──
    h2(doc, "2.4  Quizzes & Assessments")

    add_body(doc,
        "Each module concludes with a formal quiz assessment that evaluates your "
        "understanding of the material covered in that module. Passing the quiz is "
        "required to advance to the next module and ultimately to earn your certificate."
    )

    h3(doc, "When Quizzes Become Available")
    add_body(doc,
        "The quiz for each module becomes available once you have engaged with the "
        "module content. The quiz panel is accessible from the module page and will "
        "indicate when you are ready to attempt the assessment."
    )

    h3(doc, "Taking a Quiz")
    add_numbered_step(doc, 1,
        "Navigate to the module page and locate the Quiz Panel on the right side of "
        "the screen (or below the content on smaller screens)."
    )
    add_numbered_step(doc, 2, "Click the \"Start Quiz\" button to begin the assessment.")
    add_numbered_step(doc, 3,
        "Each question is presented as a multiple-choice question with several answer "
        "options. Read each question carefully and select your answer."
    )
    add_numbered_step(doc, 4,
        "After answering all questions, submit the quiz to receive your score."
    )

    add_screenshot_placeholder(doc, "Quiz Panel \u2014 Question Display with Answer Selection")

    h3(doc, "Understanding Your Score")
    add_body(doc,
        "After submitting your quiz, your score is calculated as a percentage of "
        "correct answers. A results modal will display your score and indicate whether "
        "you passed or failed."
    )
    add_body_bold_lead(doc, "Passing Threshold: ",
        "You must score 85% or higher to pass the quiz. This threshold ensures that "
        "you have a solid understanding of the material before progressing."
    )
    add_body_bold_lead(doc, "If You Pass: ",
        "A congratulatory message will appear. The next module will be unlocked "
        "automatically, and your dashboard will update to reflect your progress."
    )
    add_body_bold_lead(doc, "If You Do Not Pass: ",
        "A message will indicate that you did not reach the 85% threshold. You are "
        "encouraged to review the module content and retake the quiz."
    )

    add_callout_box(doc, "Important",
        "You have unlimited attempts to retake any quiz. There is no penalty for "
        "retaking a quiz. Your highest score is recorded.",
        "important"
    )

    add_screenshot_placeholder(doc, "Quiz Results Modal \u2014 Passed")
    add_screenshot_placeholder(doc, "Quiz Results Modal \u2014 Not Passed")

    # ── 2.5 Key Takeaways ──
    h2(doc, "2.5  Key Takeaways")

    add_body(doc,
        "After completing each module, you have the opportunity to write personal key "
        "takeaways. This feature allows you to reflect on what you learned and capture "
        "the most important points in your own words."
    )

    h3(doc, "Writing Your Takeaways")
    add_body(doc,
        "After finishing a module's content and quiz, a text area will be available for "
        "you to enter your key takeaways. Type your reflections and submit them. This "
        "is an optional but highly recommended step that reinforces learning through "
        "active recall."
    )

    add_callout_box(doc, "Note",
        "Your key takeaways are submitted to your administrator for review. As a student, "
        "you will not see a summary of your takeaways on your dashboard. Your Company "
        "Administrator or Super Administrator may use your takeaways to assess engagement "
        "and understanding across the team.",
        "note"
    )

    # ── 2.6 Chat To A Builder ──
    h2(doc, "2.6  Chat To A Builder \u2014 AI Assistant")

    add_body(doc,
        "The Build A Home LMS includes an integrated AI-powered assistant called "
        "\"Chat To A Builder.\" This chatbot is designed to help you with questions "
        "about homebuilding and the course content as you progress through the modules."
    )

    h3(doc, "What the AI Assistant Can Help With")
    add_body(doc,
        "The Chat To A Builder assistant is trained to answer questions related to "
        "homebuilding and construction. You can use it to:"
    )
    add_bullet(doc, "Ask clarifying questions about concepts covered in the modules.")
    add_bullet(doc, "Get additional explanations for topics you find challenging.")
    add_bullet(doc, "Explore construction-related topics in greater depth.")
    add_bullet(doc, "Receive expert guidance on homebuilding best practices.")

    h3(doc, "How to Use the Chatbot")
    add_numbered_step(doc, 1,
        "Look for the chat icon or \"Chat To A Builder\" button in the platform interface. "
        "Click it to open the chat window."
    )
    add_numbered_step(doc, 2,
        "The chatbot will greet you with a welcome message. Type your question into the "
        "text area at the bottom of the chat window."
    )
    add_numbered_step(doc, 3,
        "Press Enter on your keyboard to send your message, or click the Send button "
        "(arrow icon). To type a multi-line message, use Shift+Enter to create a new "
        "line without sending."
    )
    add_numbered_step(doc, 4,
        "The chatbot will process your question and respond with relevant information. "
        "The conversation is displayed in a familiar messaging format, with your messages "
        "on the right and the chatbot's responses on the left."
    )
    add_numbered_step(doc, 5,
        "When you are finished, close the chat window by clicking the X button or "
        "clicking outside the chat window."
    )

    add_callout_box(doc, "Tip",
        "For the best results, ask specific questions. Instead of \"Tell me about "
        "foundations,\" try \"What are the differences between slab-on-grade and "
        "crawl space foundations?\"",
        "tip"
    )

    add_callout_box(doc, "Note",
        "Each time you open the chat window, a new conversation session begins. "
        "Previous chat history is not retained between sessions.",
        "note"
    )

    add_screenshot_placeholder(doc, "Chat To A Builder \u2014 Chat Window with Conversation")

    # ── 2.7 Course Completion & Certification ──
    h2(doc, "2.7  Course Completion & Certification")

    add_body(doc,
        "Upon successfully completing all 10 modules and passing all quizzes, you become "
        "eligible to receive a professional certificate of completion from Strategem."
    )

    h3(doc, "Certificate Eligibility Requirements")
    add_body(doc, "To be eligible for your certificate, you must meet all of the following criteria:")
    add_bullet(doc, "Complete all 10 learning modules (view all content).")
    add_bullet(doc, "Pass all 10 module quizzes with a score of 85% or higher.")

    h3(doc, "Checking Your Eligibility")
    add_body(doc,
        "Your student dashboard will indicate when you have met all requirements for "
        "the certificate. A certificate section or button will appear on your dashboard "
        "once you are eligible."
    )

    h3(doc, "Downloading Your Certificate")
    add_body(doc,
        "Once eligible, you can download your certificate as a PDF document. Click the "
        "\"Download Certificate\" button on your dashboard. The certificate will include "
        "your full name (as entered during registration), the course title, the date of "
        "completion, and the Strategem branding."
    )

    add_callout_box(doc, "Tip",
        "Ensure that your full name was entered correctly during registration, as it "
        "will appear exactly as entered on your certificate. If your name is incorrect, "
        "contact your administrator to request a correction before downloading.",
        "tip"
    )

    add_screenshot_placeholder(doc, "Certificate Download Button on Student Dashboard")
    add_screenshot_placeholder(doc, "Sample Certificate \u2014 PDF Preview")

    # ── 2.8 End-of-Course Survey ──
    h2(doc, "2.8  End-of-Course Survey")

    add_body(doc,
        "After completing all modules and quizzes, you will be presented with an "
        "end-of-course survey. This survey helps Strategem gather feedback to improve "
        "the course and platform experience for future learners."
    )

    h3(doc, "Completing the Survey")
    add_body(doc,
        "The survey modal will appear automatically upon course completion, or you may "
        "access it from your dashboard. The survey includes:"
    )
    add_bullet(doc, "A star rating to indicate your overall satisfaction with the course.")
    add_bullet(doc, "A text field where you can provide detailed feedback, suggestions, or comments.")
    add_body(doc,
        "Complete the survey fields and click \"Submit\" to record your response. Your "
        "feedback is anonymous in the sense that it is used for aggregate analysis and "
        "course improvement purposes."
    )

    add_callout_box(doc, "Note",
        "The survey is optional but highly encouraged. Your feedback directly contributes "
        "to improving the learning experience for future students.",
        "note"
    )

    add_screenshot_placeholder(doc, "End-of-Course Survey Modal")

    # ── 2.9 Getting Help ──
    h2(doc, "2.9  Getting Help")

    add_body(doc,
        "If you encounter any issues while using the platform, or if you have questions "
        "that are not answered in this manual, several support options are available to you."
    )

    h3(doc, "Using the Help Button")
    add_body(doc,
        "The platform includes a Help button accessible from any page. Clicking this "
        "button opens a help modal with a contact form that allows you to send a message "
        "directly to the platform support team."
    )

    h3(doc, "Contact Form")
    add_body(doc, "The contact form includes the following fields:")
    add_bullet(doc, "Name \u2014 Your full name.")
    add_bullet(doc, "Email \u2014 Your email address for the support team to respond to.")
    add_bullet(doc, "Subject \u2014 A brief description of your issue or question.")
    add_bullet(doc, "Message \u2014 A detailed explanation of the issue or question.")
    add_body(doc,
        "After filling out the form, click \"Send\" to submit your request. The support "
        "team will respond to the email address you provided."
    )

    h3(doc, "Contacting Your Administrator")
    add_body(doc,
        "If you were enrolled through a company, your Company Administrator is also "
        "available to assist with account-related issues, such as company code problems "
        "or enrollment questions. Contact your administrator directly through your "
        "organization's internal communication channels."
    )

    add_screenshot_placeholder(doc, "Help Modal \u2014 Contact Form")

    add_section_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PART 3: COMPANY ADMINISTRATOR GUIDE
    # ═══════════════════════════════════════════════════════════════════════

    h1(doc, "Part 3: Company Administrator Guide")
    add_body(doc,
        "This section is intended for Company Administrators \u2014 users who manage students "
        "within their organization on the Build A Home LMS platform. As a Company "
        "Administrator, you have access to tools for monitoring student progress, viewing "
        "engagement data, and exporting reports."
    )

    # ── 3.1 Getting Started ──
    h2(doc, "3.1  Getting Started")

    h3(doc, "Accessing the Platform")
    add_body(doc,
        "Access the Build A Home LMS by navigating to the platform URL in your web "
        "browser. Use the same URL that your students use to access the platform."
    )

    h3(doc, "Signing In")
    add_body(doc,
        "Sign in using the email and password associated with your Company Administrator "
        "account. On the sign-in page, enter your credentials under the \"Login\" tab and "
        "click \"Sign In.\" Upon successful authentication, you will be directed to the "
        "Company Administrator dashboard."
    )

    add_callout_box(doc, "Note",
        "Company Administrator accounts are assigned by a Super Administrator. If you "
        "do not have an account or need your role changed, contact your Super Administrator.",
        "note"
    )

    h3(doc, "Understanding Your Role and Permissions")
    add_body(doc,
        "As a Company Administrator, you have the following capabilities:"
    )
    add_bullet(doc, "View aggregate statistics for your company's enrolled students.")
    add_bullet(doc, "Monitor individual student progress and quiz scores.")
    add_bullet(doc, "View key takeaways submitted by your students.")
    add_bullet(doc, "Export student data for reporting purposes.")
    add_body(doc,
        "You do not have access to platform-wide settings, user management for other "
        "companies, or module configuration. These capabilities are reserved for Super "
        "Administrators."
    )

    # ── 3.2 Company Admin Dashboard ──
    h2(doc, "3.2  Company Admin Dashboard")

    h3(doc, "Dashboard Layout")
    add_body(doc,
        "The Company Admin Dashboard is your primary workspace. It provides an at-a-glance "
        "view of your organization's engagement with the course."
    )
    add_body_bold_lead(doc, "Company Statistics Cards: ",
        "At the top of the dashboard, summary cards display key metrics including the "
        "total number of enrolled students, the average completion rate across your "
        "students, and the number of students who have completed the full course."
    )
    add_body_bold_lead(doc, "Student Roster: ",
        "Below the statistics, a table lists all students associated with your company. "
        "Each row displays the student's name, email, progress percentage, number of "
        "modules completed, and quiz pass status."
    )

    add_screenshot_placeholder(doc, "Company Admin Dashboard \u2014 Statistics Cards and Student Roster")

    h3(doc, "Key Metrics")
    add_body(doc, "The dashboard tracks the following key performance indicators:")
    add_bullet(doc, "Total Students \u2014 The number of students registered with your company code.")
    add_bullet(doc, "Completion Rate \u2014 The percentage of your students who have completed all 10 modules.")
    add_bullet(doc, "Average Progress \u2014 The mean progress percentage across all of your enrolled students.")

    # ── 3.3 Managing Your Students ──
    h2(doc, "3.3  Managing Your Students")

    h3(doc, "Viewing Your Student Roster")
    add_body(doc,
        "The student roster on your dashboard lists every student who registered with "
        "your organization's company code. For each student, you can view:"
    )
    add_bullet(doc, "Full name and email address.")
    add_bullet(doc, "Overall course progress (percentage).")
    add_bullet(doc, "The number of modules completed out of 10.")
    add_bullet(doc, "Individual quiz scores for each module attempted.")

    h3(doc, "Student Progress Details")
    add_body(doc,
        "To view detailed progress information for a specific student, locate their row "
        "in the roster table. The progress column shows a visual progress bar along "
        "with the exact percentage. You can identify students who may be falling behind "
        "or who have completed the course and are ready for certification."
    )

    h3(doc, "Filtering and Searching")
    add_body(doc,
        "If your organization has a large number of enrolled students, you can use the "
        "search and filter tools available on the roster to quickly locate specific "
        "students by name or email, or to filter by completion status."
    )

    add_screenshot_placeholder(doc, "Student Roster \u2014 Search and Filter Options")

    # ── 3.4 Viewing Student Takeaways ──
    h2(doc, "3.4  Viewing Student Takeaways")

    add_body(doc,
        "Students are encouraged to write personal key takeaways after completing each "
        "module. As a Company Administrator, you have the ability to review the takeaways "
        "submitted by students within your organization."
    )
    add_body(doc,
        "This feature provides valuable insight into how your employees are engaging "
        "with the material and what concepts are resonating most with your team. Use "
        "this information to identify areas where additional support or discussion may "
        "be beneficial."
    )

    add_screenshot_placeholder(doc, "Student Takeaways View \u2014 Company Admin")

    # ── 3.5 Company Codes ──
    h2(doc, "3.5  Company Codes")

    h3(doc, "What Are Company Codes?")
    add_body(doc,
        "Company codes are unique identifiers that link students to your organization "
        "during registration. When a student enters your company code while creating "
        "their account, they are automatically associated with your company, and their "
        "progress will appear on your Company Admin Dashboard."
    )

    h3(doc, "Distributing Codes to Employees")
    add_body(doc,
        "Your company code is assigned by a Super Administrator. Once you have received "
        "your code, distribute it to your employees along with the platform URL and "
        "registration instructions. Employees should enter the code in the "
        "\"Company Code\" field during account creation."
    )

    add_callout_box(doc, "Important",
        "Company codes must be entered at the time of registration. Students cannot "
        "add or change their company code after account creation. Ensure your employees "
        "have the correct code before they register.",
        "important"
    )

    # ── 3.6 Exporting Data ──
    h2(doc, "3.6  Exporting Data")

    h3(doc, "Exporting Student Data")
    add_body(doc,
        "The platform allows you to export your company's student data as a CSV file "
        "for use in external reporting tools, spreadsheets, or HR systems. To export:"
    )
    add_numbered_step(doc, 1, "Navigate to your Company Admin Dashboard.")
    add_numbered_step(doc, 2, "Locate the export option (typically an \"Export\" or \"Download\" button).")
    add_numbered_step(doc, 3, "Click the button to download a CSV file containing student records.")

    h3(doc, "Understanding Export Fields")
    add_body(doc, "The exported data typically includes the following fields:")
    add_bullet(doc, "Student name and email address.")
    add_bullet(doc, "Registration date.")
    add_bullet(doc, "Overall progress percentage.")
    add_bullet(doc, "Number of modules completed.")
    add_bullet(doc, "Individual quiz scores.")
    add_bullet(doc, "Course completion status.")

    add_screenshot_placeholder(doc, "Export Button and Sample CSV Output")

    # ── 3.7 Getting Help ──
    h2(doc, "3.7  Getting Help")

    add_body(doc,
        "If you need assistance with your Company Administrator account or have "
        "questions about the platform, the following support options are available:"
    )
    add_body_bold_lead(doc, "Help Button: ",
        "Use the in-platform Help button to access the contact form and send a message "
        "to the Strategem support team."
    )
    add_body_bold_lead(doc, "Super Administrator: ",
        "For account permissions, role changes, or company code issues, contact your "
        "organization's Super Administrator."
    )
    add_body_bold_lead(doc, "Strategem Support: ",
        "For technical issues or platform-related questions, submit a support request "
        "through the Help contact form."
    )

    add_section_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PART 4: SUPER ADMINISTRATOR GUIDE
    # ═══════════════════════════════════════════════════════════════════════

    h1(doc, "Part 4: Super Administrator Guide")
    add_body(doc,
        "This section is intended for Super Administrators \u2014 users with full access to "
        "all platform management features. As a Super Administrator, you can manage users, "
        "companies, modules, company codes, analytics, data exports, and more. This guide "
        "covers every administrative function available to you."
    )

    # ── 4.1 Getting Started ──
    h2(doc, "4.1  Getting Started")

    h3(doc, "Accessing the Platform")
    add_body(doc,
        "Access the Build A Home LMS by navigating to the platform URL in your web "
        "browser. Super Administrators use the same sign-in page as all other users."
    )

    h3(doc, "Signing In")
    add_body(doc,
        "Sign in using the email and password associated with your Super Administrator "
        "account. Upon successful authentication, you will be directed to the Super "
        "Administrator dashboard, which provides access to all platform management tools."
    )

    h3(doc, "Super Admin Dashboard Overview")
    add_body(doc,
        "The Super Admin Dashboard serves as your command center for the entire platform. "
        "It is divided into several functional areas:"
    )
    add_bullet(doc, "Platform-wide statistics and analytics at the top of the dashboard.")
    add_bullet(doc, "Quick-access navigation to major management areas (Users, Companies, Modules, etc.).")
    add_bullet(doc, "Recent activity feeds and summary tables for an at-a-glance view of platform health.")

    add_screenshot_placeholder(doc, "Super Admin Dashboard \u2014 Full Overview")

    # ── 4.2 Platform Analytics ──
    h2(doc, "4.2  Platform Analytics")

    add_body(doc,
        "The Super Admin Dashboard provides access to comprehensive analytics that "
        "help you understand how the platform is being used across all companies and users."
    )

    h3(doc, "Platform-Wide Statistics")
    add_body(doc,
        "At the top of your dashboard, summary cards display key metrics including:"
    )
    add_bullet(doc, "Total registered users across all companies.")
    add_bullet(doc, "Total active students currently progressing through the course.")
    add_bullet(doc, "Overall completion rate across the entire platform.")
    add_bullet(doc, "Total number of companies enrolled.")

    h3(doc, "User Statistics and Trends")
    add_body(doc,
        "The user statistics section provides detailed breakdowns of user activity, "
        "including registration trends, active user counts, and engagement patterns "
        "over time."
    )

    h3(doc, "Module Analytics")
    add_body(doc,
        "Module-level analytics allow you to see how students are performing on each "
        "individual module. Key metrics include:"
    )
    add_bullet(doc, "Completion rate per module \u2014 the percentage of students who have finished each module.")
    add_bullet(doc, "Average quiz score per module \u2014 the mean score achieved by students on each quiz.")
    add_bullet(doc, "Drop-off analysis \u2014 identify modules where students are most likely to stop progressing.")

    add_screenshot_placeholder(doc, "Platform Analytics \u2014 Statistics Cards and Module Performance")

    # ── 4.3 User Management ──
    h2(doc, "4.3  User Management")

    add_body(doc,
        "The User Management section allows you to view, search, and manage all "
        "registered users on the platform, regardless of their company affiliation."
    )

    h3(doc, "Viewing All Users")
    add_body(doc,
        "The user management interface displays a searchable, sortable table of all "
        "platform users. For each user, you can view their name, email, role, company "
        "affiliation, registration date, and course progress."
    )
    add_body(doc,
        "Use the search bar to find users by name or email. Filter options allow you to "
        "narrow the list by role (Student, Company Admin, Super Admin) or by company."
    )

    add_screenshot_placeholder(doc, "User Management \u2014 User List with Search and Filters")

    h3(doc, "Changing User Roles")
    add_body(doc,
        "To change a user's role (for example, to promote a student to Company Admin):"
    )
    add_numbered_step(doc, 1, "Locate the user in the user management table.")
    add_numbered_step(doc, 2, "Click on the user's record to view their details.")
    add_numbered_step(doc, 3, "Select the new role from the role assignment options (Student, Company Admin, or Super Admin).")
    add_numbered_step(doc, 4, "Confirm the change. The user's permissions will be updated immediately.")

    add_callout_box(doc, "Warning",
        "Exercise caution when assigning the Super Administrator role. Super Admins "
        "have unrestricted access to all platform data and management functions. Only "
        "assign this role to trusted personnel.",
        "warning"
    )

    h3(doc, "Assigning Users to Companies")
    add_body(doc,
        "You can manually assign or reassign a user to a company. This is useful when "
        "a user registered without a company code or needs to be transferred to a "
        "different organization:"
    )
    add_numbered_step(doc, 1, "Locate the user in the user management table.")
    add_numbered_step(doc, 2, "Click on the user's record to access their details.")
    add_numbered_step(doc, 3, "Select the target company from the company assignment dropdown.")
    add_numbered_step(doc, 4, "Save the change. The user will now appear under the selected company's roster.")

    h3(doc, "Deleting Users")
    add_body(doc,
        "To remove a user from the platform permanently:"
    )
    add_numbered_step(doc, 1, "Locate the user in the user management table.")
    add_numbered_step(doc, 2, "Click on the user's record and select the delete option.")
    add_numbered_step(doc, 3, "Confirm the deletion when prompted.")

    add_callout_box(doc, "Important",
        "Deleting a user is permanent and cannot be undone. All of the user's progress "
        "data, quiz scores, and takeaways will be removed from the system. Ensure you "
        "have exported any necessary data before proceeding.",
        "important"
    )

    h3(doc, "Purging Inactive Accounts")
    add_body(doc,
        "The platform includes a bulk action to purge inactive accounts \u2014 users who "
        "registered but have not engaged with the course content beyond a defined "
        "threshold. This feature helps keep the user database clean and accurate."
    )
    add_body(doc,
        "To purge inactive accounts, navigate to the purge option in the user management "
        "area. Review the list of accounts flagged as inactive, confirm the action, and "
        "the selected accounts will be permanently removed."
    )

    add_callout_box(doc, "Warning",
        "Purging is a bulk deletion operation and cannot be undone. Carefully review "
        "the list of affected accounts before confirming.",
        "warning"
    )

    add_screenshot_placeholder(doc, "User Management \u2014 Purge Inactive Accounts Confirmation")

    # ── 4.4 Company Management ──
    h2(doc, "4.4  Company Management")

    add_body(doc,
        "The Company Management section allows you to create, edit, and manage the "
        "organizations that are enrolled on the platform."
    )

    h3(doc, "Viewing All Companies")
    add_body(doc,
        "The company management interface displays a list of all registered companies "
        "with key information including company name, number of enrolled students, "
        "overall completion rate, and active status."
    )

    add_screenshot_placeholder(doc, "Company Management \u2014 Company List")

    h3(doc, "Creating a New Company")
    add_numbered_step(doc, 1, "Navigate to the Company Management section.")
    add_numbered_step(doc, 2, "Click the \"Add Company\" or \"Create Company\" button.")
    add_numbered_step(doc, 3, "Enter the company name and any required details in the form.")
    add_numbered_step(doc, 4, "Save the new company. It will now appear in the company list and can be assigned a company code.")

    h3(doc, "Editing Company Details")
    add_body(doc,
        "To edit an existing company's information, click on the company in the list to "
        "open its detail view. Update the necessary fields and save your changes."
    )

    h3(doc, "Deactivating a Company")
    add_body(doc,
        "If a company is no longer enrolled in the program, you can deactivate it rather "
        "than deleting it. Deactivating a company preserves its historical data (student "
        "records, progress, and scores) while preventing new registrations with that "
        "company's code."
    )

    h3(doc, "Company-Specific Statistics")
    add_body(doc,
        "From the company detail view, you can access statistics specific to that "
        "company, including the number of enrolled students, average progress, completion "
        "rates, and individual student performance data."
    )

    add_screenshot_placeholder(doc, "Company Detail View \u2014 Statistics and Student List")

    # ── 4.5 Company Code Management ──
    h2(doc, "4.5  Company Code Management")

    add_body(doc,
        "Company codes are the mechanism by which students are associated with their "
        "employer or organization during registration. The Company Code Management "
        "section allows you to create, assign, and manage these codes."
    )

    h3(doc, "Viewing All Codes")
    add_body(doc,
        "The code management interface displays a list of all company codes, along with "
        "their associated company, status (active or inactive), and usage count."
    )

    h3(doc, "Generating New Codes")
    add_numbered_step(doc, 1, "Navigate to the Company Code Management section.")
    add_numbered_step(doc, 2, "Click the \"Generate Code\" or \"Add Code\" button.")
    add_numbered_step(doc, 3, "Select the company to associate the code with.")
    add_numbered_step(doc, 4, "Configure any additional settings (such as expiration or usage limits, if applicable).")
    add_numbered_step(doc, 5, "Save the code. It is now ready to be distributed to the company's employees.")

    h3(doc, "Activating and Deactivating Codes")
    add_body(doc,
        "You can activate or deactivate codes at any time. Deactivating a code prevents "
        "new students from registering with it, but does not affect students who have "
        "already registered using that code."
    )

    add_screenshot_placeholder(doc, "Company Code Management \u2014 Code List with Actions")

    # ── 4.6 Viewing User Takeaways ──
    h2(doc, "4.6  Viewing User Takeaways")

    add_body(doc,
        "As a Super Administrator, you have access to all student takeaways submitted "
        "across the entire platform. This provides a comprehensive view of student "
        "engagement and understanding across all companies and modules."
    )
    add_body(doc,
        "The takeaways management interface allows you to filter by user, module, or "
        "company to identify trends, assess engagement levels, and evaluate how "
        "effectively the course content is being absorbed by students."
    )
    add_body(doc,
        "Use this data to inform decisions about content updates, identify modules "
        "that may need enhancement, or recognize companies and students that are "
        "particularly engaged with the material."
    )

    add_screenshot_placeholder(doc, "User Takeaways \u2014 Super Admin View with Filters")

    # ── 4.7 Module Management ──
    h2(doc, "4.7  Module Management")

    add_body(doc,
        "The Module Management section provides tools to create, edit, organize, and "
        "control the course modules that students interact with."
    )

    h3(doc, "Viewing All Modules")
    add_body(doc,
        "The module management interface displays all course modules in their current "
        "order, along with their title, status (active or inactive), and key statistics "
        "such as completion rates and average quiz scores."
    )

    add_screenshot_placeholder(doc, "Module Management \u2014 Module List")

    h3(doc, "Creating a New Module")
    add_numbered_step(doc, 1, "Navigate to the Module Management section.")
    add_numbered_step(doc, 2, "Click the \"Add Module\" or \"Create Module\" button.")
    add_numbered_step(doc, 3, "Enter the module title, description, and content.")
    add_numbered_step(doc, 4, "Configure the module's video content URL and quiz questions.")
    add_numbered_step(doc, 5, "Set the module's position in the sequence order.")
    add_numbered_step(doc, 6, "Save the module. It will be available to students based on its active status and sequence position.")

    h3(doc, "Editing Module Content and Settings")
    add_body(doc,
        "To edit an existing module, click on it in the module list to open the editing "
        "interface. You can update the module title, description, content, video, quiz "
        "questions, and other settings. Save your changes when finished."
    )

    add_callout_box(doc, "Important",
        "Changes to module content are reflected immediately for all users. If students "
        "are currently progressing through a module you are editing, consider timing "
        "your updates to minimize disruption.",
        "important"
    )

    h3(doc, "Reordering Modules")
    add_body(doc,
        "The course sequence is determined by the order of modules in the management "
        "interface. You can reorder modules by adjusting their sequence position. "
        "Students will encounter modules in the order you define."
    )

    h3(doc, "Activating and Deactivating Modules")
    add_body(doc,
        "You can activate or deactivate individual modules. Deactivating a module "
        "removes it from the student-facing course without deleting it. This is useful "
        "for temporarily removing content that is being revised or for phased course "
        "rollouts."
    )

    # ── 4.8 Survey Results ──
    h2(doc, "4.8  Survey Results")

    add_body(doc,
        "The Survey Results section provides access to all feedback submitted by "
        "students through the end-of-course survey. This data is valuable for "
        "understanding student satisfaction and identifying areas for improvement."
    )

    h3(doc, "Viewing Survey Data")
    add_body(doc,
        "The survey results interface displays aggregated data including average "
        "ratings, total number of responses, and individual text feedback from students. "
        "You can review this data to identify common themes, praises, and constructive "
        "criticism."
    )

    h3(doc, "Filtering Results")
    add_body(doc,
        "Filter survey results by date range or company to analyze feedback from "
        "specific time periods or organizational groups. This is particularly useful "
        "after onboarding a new company cohort to assess their experience."
    )

    add_screenshot_placeholder(doc, "Survey Results \u2014 Aggregated Data and Individual Responses")

    # ── 4.9 Data Exports ──
    h2(doc, "4.9  Data Exports")

    add_body(doc,
        "The platform provides comprehensive data export capabilities for reporting, "
        "auditing, and analysis purposes. As a Super Administrator, you can export data "
        "at multiple levels."
    )

    h3(doc, "Exporting User Data")
    add_body(doc,
        "Export a complete list of all registered users, including their names, email "
        "addresses, roles, company affiliations, registration dates, and current "
        "progress. The data is downloaded as a CSV file."
    )

    h3(doc, "Exporting Company Data")
    add_body(doc,
        "Export a summary of all companies on the platform, including company names, "
        "enrollment counts, completion rates, and other aggregate statistics."
    )

    h3(doc, "Exporting Student Progress Data")
    add_body(doc,
        "Export detailed student progress data, including per-module completion status, "
        "quiz scores, and overall course progress. This export can be filtered by "
        "company for targeted reporting."
    )

    add_callout_box(doc, "Tip",
        "Schedule regular data exports for your records. Exported CSV files can be "
        "opened in Microsoft Excel, Google Sheets, or any spreadsheet application for "
        "further analysis.",
        "tip"
    )

    add_screenshot_placeholder(doc, "Data Export Options \u2014 Super Admin")

    # ── 4.10 User Emulation ──
    h2(doc, "4.10  User Emulation")

    add_body(doc,
        "User Emulation is a powerful diagnostic tool that allows you to view the "
        "platform exactly as a specific user sees it. This is invaluable for "
        "troubleshooting user-reported issues, verifying progress data, and testing "
        "the student or company admin experience."
    )

    h3(doc, "What Emulation Is and When to Use It")
    add_body(doc,
        "When you emulate a user, you temporarily adopt their identity within the "
        "platform. You will see their dashboard, their progress, their modules, and "
        "all data as if you were logged in as that user. Emulation is read-only in the "
        "sense that it is intended for observation and diagnostics, not for making changes "
        "on behalf of the user."
    )
    add_body(doc, "Common use cases for emulation include:")
    add_bullet(doc, "Verifying a student's reported issue with progress or quiz scores.")
    add_bullet(doc, "Reviewing a student's dashboard layout and content display.")
    add_bullet(doc, "Testing the experience of a Company Administrator for a specific company.")

    h3(doc, "Starting an Emulation Session")
    add_numbered_step(doc, 1, "Navigate to the User Management section.")
    add_numbered_step(doc, 2, "Locate the user you wish to emulate.")
    add_numbered_step(doc, 3, "Click the \"Emulate\" button or option associated with that user.")
    add_numbered_step(doc, 4,
        "The platform will reload and display the user's view. A banner or indicator "
        "will confirm that you are in emulation mode."
    )

    h3(doc, "Returning to Your Admin Account")
    add_body(doc,
        "To exit emulation mode and return to your Super Administrator account, "
        "click the \"Stop Emulating\" or \"Return to Admin\" button that appears in the "
        "emulation indicator banner. You will be immediately returned to your Super "
        "Admin dashboard."
    )

    add_callout_box(doc, "Warning",
        "While emulating a user, be careful not to perform actions that could affect "
        "the user's data. Always exit emulation mode when you have finished your review.",
        "warning"
    )

    add_screenshot_placeholder(doc, "User Emulation \u2014 Emulation Mode Indicator Banner")

    # ── 4.11 Getting Help ──
    h2(doc, "4.11  Getting Help")

    add_body(doc,
        "As a Super Administrator, you have the highest level of platform access. If you "
        "encounter technical issues that you cannot resolve through the administrative "
        "tools, the following support options are available:"
    )
    add_body_bold_lead(doc, "Help Button: ",
        "Use the in-platform Help button to submit a support request to the Strategem "
        "technical team."
    )
    add_body_bold_lead(doc, "Strategem Support: ",
        "For critical platform issues, contact Strategem support directly through the "
        "help contact form or through your organization's support channel."
    )

    add_section_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PART 5: TROUBLESHOOTING & FAQ
    # ═══════════════════════════════════════════════════════════════════════

    h1(doc, "Part 5: Troubleshooting & FAQ")
    add_body(doc,
        "This section addresses common issues that users may encounter while using "
        "the Build A Home LMS platform and provides solutions and guidance for "
        "resolving them."
    )

    h2(doc, "Common Login Issues")

    add_body_bold_lead(doc, "I cannot sign in with my email and password. ",
        "Verify that you are entering the correct email address (the one used during "
        "registration) and that your password is typed correctly, including any uppercase "
        "or special characters. If you have forgotten your password, use the "
        "\"Forgot Password?\" link on the sign-in page to reset it."
    )
    add_body_bold_lead(doc, "I never received a password reset email. ",
        "Check your spam or junk folder. If the email is not there, wait a few minutes "
        "and try again. Ensure you are entering the email address associated with your "
        "account. If the issue persists, contact support through the Help form."
    )
    add_body_bold_lead(doc, "My account does not exist or has been deactivated. ",
        "If you receive a message indicating your account does not exist, you may need "
        "to create a new account. If you believe your account was deactivated in error, "
        "contact your Company Administrator or the platform support team."
    )

    h2(doc, "Quiz and Progress Issues")

    add_body_bold_lead(doc, "My quiz score seems incorrect. ",
        "Quiz scores are calculated automatically based on the number of correct answers. "
        "Review your answers carefully. If you believe there is an error, contact support "
        "with the module name and your quiz details."
    )
    add_body_bold_lead(doc, "I passed the quiz but the next module is still locked. ",
        "Try refreshing your browser. If the issue persists, sign out and sign back in. "
        "Your progress is saved on the server, so logging in again should refresh your "
        "dashboard correctly."
    )
    add_body_bold_lead(doc, "My progress does not seem to be saving. ",
        "Ensure you have a stable internet connection. Progress is saved periodically "
        "as you engage with the content. If you are experiencing connectivity issues, "
        "your progress may not be recorded until the connection is restored."
    )

    h2(doc, "Certificate Issues")

    add_body_bold_lead(doc, "I completed all modules but cannot download my certificate. ",
        "Verify that you have passed all 10 module quizzes with a score of 85% or higher. "
        "Even if you completed the module content, the certificate requires a passing quiz "
        "score for every module. Check your dashboard for any modules that show as incomplete."
    )
    add_body_bold_lead(doc, "My name is incorrect on the certificate. ",
        "The certificate uses the name you entered during registration. Contact your "
        "Company Administrator or Super Administrator to request a name correction in "
        "the system, then download the certificate again."
    )

    h2(doc, "Browser Compatibility")

    add_body_bold_lead(doc, "The platform does not display correctly. ",
        "Ensure you are using a supported browser (Chrome, Firefox, Edge, or Safari) "
        "and that it is updated to the latest version. Clear your browser cache and "
        "cookies, then reload the page. Disable any browser extensions that may interfere "
        "with the platform (ad blockers, script blockers, etc.)."
    )
    add_body_bold_lead(doc, "Videos are not playing. ",
        "Check that your browser supports HTML5 video playback. Ensure you are not on "
        "a restricted network that blocks video streaming. Try a different browser or "
        "device to isolate the issue."
    )

    h2(doc, "Contact Support")

    add_body(doc,
        "If your issue is not addressed above, or if you have tried the suggested "
        "solutions without success, please contact the Strategem support team using "
        "the in-platform Help button. Provide as much detail as possible, including:"
    )
    add_bullet(doc, "Your name and email address.")
    add_bullet(doc, "A description of the issue and when it occurred.")
    add_bullet(doc, "The browser and device you are using.")
    add_bullet(doc, "Any error messages you received (screenshots are helpful).")

    add_section_break(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # APPENDIX: GLOSSARY
    # ═══════════════════════════════════════════════════════════════════════

    h1(doc, "Appendix: Glossary of Terms")
    doc.add_paragraph()

    terms = [
        ("Build A Home", "The name of the 10-module construction training course delivered through the Strategem LMS platform."),
        ("Certificate of Completion", "A PDF document issued to students who have completed all 10 modules and passed all quizzes with a score of 85% or higher."),
        ("Chat To A Builder", "The platform's integrated AI-powered chatbot that answers questions about homebuilding and course content."),
        ("Company Administrator", "A user role with permissions to view and manage students within a specific company or organization."),
        ("Company Code", "A unique code assigned to an organization that students enter during registration to be associated with that company."),
        ("Dashboard", "The main landing page after sign-in, displaying progress, statistics, and navigation options specific to the user's role."),
        ("Emulation", "A Super Administrator feature that allows viewing the platform as if logged in as another user, for diagnostic purposes."),
        ("Knowledge Check", "An interactive self-assessment question embedded within module content. Does not count toward quiz scores."),
        ("LMS", "Learning Management System \u2014 a software platform for delivering, tracking, and managing educational content."),
        ("Module", "A unit of course content consisting of video instruction, written material, knowledge checks, and a quiz assessment."),
        ("Module Map", "A table of contents within a module that provides clickable links to each section of the module's content."),
        ("Quiz", "A formal assessment at the end of each module. Students must score 85% or higher to pass and advance."),
        ("Strategem", "Strategem LLC, the company that developed and operates the Build A Home LMS platform."),
        ("Student", "A user role for individuals enrolled in the course who are completing modules and assessments."),
        ("Super Administrator", "A user role with full, unrestricted access to all platform management features."),
        ("Takeaway", "A personal reflection or note written by a student after completing a module, submitted for administrator review."),
    ]

    # Build glossary as a table
    table = doc.add_table(rows=len(terms) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(["Term", "Definition"]):
        cell = table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, "F97316")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)

    # Data rows
    for r, (term, definition) in enumerate(terms):
        cell_term = table.cell(r + 1, 0)
        cell_term.text = term
        for p in cell_term.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = DARK_BLUE

        cell_def = table.cell(r + 1, 1)
        cell_def.text = definition
        for p in cell_def.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = DARK_BLUE

        if r % 2 == 1:
            set_cell_shading(cell_term, "FFF7ED")
            set_cell_shading(cell_def, "FFF7ED")

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Footer notice ──
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "\u00A9 " + str(date.today().year) + " Strategem LLC. All rights reserved.\n"
        "This document is confidential and intended for authorized users only."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_500

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"\nUser manual generated successfully!")
    print(f"Output: {OUTPUT_PATH}")
    print(f"Size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")


if __name__ == "__main__":
    build_manual()
